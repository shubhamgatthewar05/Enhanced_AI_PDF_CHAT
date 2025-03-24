import streamlit as st
from PyPDF2 import PdfReader
from langchain.text_splitter import CharacterTextSplitter, RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS, Chroma
from langchain.memory import ConversationBufferMemory
from langchain.chains import ConversationalRetrievalChain
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_openai import ChatOpenAI
from langchain.retrievers import ContextualCompressionRetriever
from langchain.retrievers.document_compressors import LLMChainExtractor
import google.generativeai as genai
import os
import pytesseract
from pdf2image import convert_from_bytes
from PIL import Image
import io
import numpy as np
import pandas as pd
import re
import tempfile
import docx
import base64
from io import BytesIO
import matplotlib.pyplot as plt
import time
import zipfile
from langchain_core.messages import HumanMessage, AIMessage
import plotly.express as px

css = '''
<style>
.chat-message {
    padding: 1.5rem; 
    border-radius: 0.8rem; 
    margin-bottom: 1rem; 
    display: flex;
    box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
}
.chat-message.user {
    background-color: #2b313e
}
.chat-message.bot {
    background-color: #475063
}
.chat-message .avatar {
  width: 15%;
}
.chat-message .avatar img {
  max-width: 78px;
  max-height: 78px;
  border-radius: 50%;
  object-fit: cover;
  border: 2px solid #fff;
}
.chat-message .message {
  width: 85%;
  padding: 0 1.5rem;
  color: #fff;
  line-height: 1.6;
}
.stExpander {
    border-radius: 8px;
    box-shadow: 0 2px 5px rgba(0,0,0,0.1);
}
.feedback-button {
    margin-right: 10px;
    padding: 5px 10px;
    border-radius: 5px;
}
.pdf-download {
    display: inline-block;
    padding: 8px 16px;
    background-color: #4CAF50;
    color: white;
    text-decoration: none;
    border-radius: 4px;
    margin: 10px 0;
}
/* Custom tab styling */
.stTabs [data-baseweb="tab-list"] {
    gap: 2px;
}
.stTabs [data-baseweb="tab"] {
    height: 50px;
    white-space: pre-wrap;
    background-color: #f0f2f6;
    border-radius: 4px 4px 0 0;
    border-right: 1px solid #e0e0e0;
    border-left: 1px solid #e0e0e0;
    border-top: 1px solid #e0e0e0;
}
.stTabs [aria-selected="true"] {
    background-color: #ffffff;
    border-bottom: 2px solid #4CAF50;
}
</style>
'''

bot_template = '''
<div class="chat-message bot">
    <div class="avatar">
        <img src="https://i.ibb.co/cN0nmSj/Screenshot-2023-05-28-at-02-37-21.png">
    </div>
    <div class="message">{{MSG}}</div>
</div>
'''

user_template = '''
<div class="chat-message user">
    <div class="avatar">
        <img src="https://i.ibb.co/rdZC7LZ/Photo-logo-1.png">
    </div>    
    <div class="message">{{MSG}}</div>
</div>
'''

def parse_file(file):
    """Parse different file types and extract text."""
    try:
        if file.type == "application/pdf":
            return parse_pdf(file)
        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return parse_docx(file)
        elif file.type == "text/plain":
            return parse_txt(file)
        elif file.type in ["application/zip", "application/x-zip-compressed"]:
            return parse_zip(file)
        else:
            st.warning(f"Unsupported file type: {file.type} for {file.name}")
            return ""
    except Exception as e:
        st.error(f"Error parsing {file.name}: {str(e)}")
        return ""

def parse_pdf(file):
    """Extract text from PDF files using PyPDF2 and OCR if needed."""
    pdf_content = file.read()

    try:
        pdf_reader = PdfReader(io.BytesIO(pdf_content))
        page_texts = []
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text.strip():
                page_texts.append(page_text)
        

        if not page_texts:
            st.info(f"No text found in {file.name} using direct extraction. Attempting OCR...")
            images = convert_from_bytes(pdf_content)
            for image in images:
                rotated_image = correct_image_orientation(image)
                page_text = extract_text_from_image(rotated_image)
                if page_text.strip():
                    page_texts.append(page_text)
        
        return "\n".join(page_texts)
    except Exception as e:
        st.warning(f"Error processing {file.name}: {str(e)}")
        return ""

def parse_docx(file):
    """Extract text from DOCX files."""
    temp = tempfile.NamedTemporaryFile(delete=False)
    temp.write(file.read())
    temp.close()
    
    doc = docx.Document(temp.name)
    text = [para.text for para in doc.paragraphs]
    os.unlink(temp.name)
    return '\n'.join(text)

def parse_txt(file):
    """Extract text from TXT files."""
    return file.read().decode('utf-8')

def parse_zip(file):
    """Extract text from all supported files in a ZIP archive."""
    combined_text = ""
    temp = tempfile.NamedTemporaryFile(delete=False)
    temp.write(file.read())
    temp.close()
    
    with zipfile.ZipFile(temp.name, 'r') as zip_ref:
        # Creating temp directory for stored files ..................
        with tempfile.TemporaryDirectory() as temp_dir:
            zip_ref.extractall(temp_dir)
            
            for root, _, files in os.walk(temp_dir):
                for filename in files:
                    file_path = os.path.join(root, filename)
                    ext = os.path.splitext(filename)[1].lower()
                    
                   
                    if ext == '.pdf':
                        with open(file_path, 'rb') as f:
                            combined_text += parse_pdf_file(f) + "\n\n"
                    elif ext == '.docx':
                        combined_text += parse_docx_file(file_path) + "\n\n"
                    elif ext == '.txt':
                        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                            combined_text += f.read() + "\n\n"
    
    os.unlink(temp.name)
    return combined_text

def parse_pdf_file(file_path):
    """Parse PDF from file path."""
    try:
        pdf_reader = PdfReader(file_path)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        return f"Error parsing PDF: {str(e)}"

def parse_docx_file(file_path):
    """Parse DOCX from file path."""
    try:
        doc = docx.Document(file_path)
        text = [para.text for para in doc.paragraphs]
        return "\n".join(text)
    except Exception as e:
        return f"Error parsing DOCX: {str(e)}"

def extract_text_from_image(image):
    try:
        # Pre-process image for better OCR results
        # Convert image to grayscale
        gray_image = image.convert('L')
        
        # Apply thresholding to binarize the image
        threshold = 128
        thresholded_image = gray_image.point(lambda x: 0 if x < threshold else 255, '1')
        
        # Optional: Increase image size/resolution for better OCR
        width, height = thresholded_image.size
        enlarged = thresholded_image.resize((width*2, height*2), Image.LANCZOS)
        
       
        custom_config = r'--oem 3 --psm 6 -l eng+fra+deu+spa -c textord_min_linesize=3'
        text = pytesseract.image_to_string(enlarged, config=custom_config)
        return text
    except Exception as e:
        st.error(f"Error in OCR processing: {str(e)}")
        return ""

def correct_image_orientation(image):
    """Correct the orientation of the image if it is rotated."""
    try:
       
        osd = pytesseract.image_to_osd(image)
        rotation = int(re.search(r'Rotate:\s*(\d+)', osd).group(1))
        
        if rotation == 90:
            return image.rotate(-90, expand=True)
        elif rotation == 180:
            return image.rotate(180, expand=True)
        elif rotation == 270:
            return image.rotate(90, expand=True)
        
        return image  
    except Exception:
        return image

def get_text_chunks(text, method="recursive"):
    if text is None or not text.strip():
        return None
    try:
        if method == "recursive":
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200,
                length_function=len,
                separators=["\n\n", "\n", " ", ""]
            )
        else:  
            text_splitter = CharacterTextSplitter(
                separator="\n",
                chunk_size=1000,
                chunk_overlap=200,
                length_function=len
            )
            
        chunks = text_splitter.split_text(text)
        if not chunks:
            st.error("No text chunks were created")
            return None
            
       
        st.session_state.chunk_info = {
            "method": method,
            "chunk_size": 1000,
            "chunk_overlap": 200,
            "total_chunks": len(chunks)
        }
        
        return chunks
    except Exception as e:
        st.error(f"Error splitting text: {str(e)}")
        return None

def get_vectorstore(text_chunks, vectorstore_type="faiss"):
    if text_chunks is None:
        return None
    try:
        with st.spinner(f"Creating {vectorstore_type.upper()} vector store..."):
          
            embeddings = HuggingFaceEmbeddings(
                model_name=st.session_state.get("embedding_model", "all-MiniLM-L6-v2"),
                model_kwargs={'device': 'cpu'}
            )
            
            # Create appropriate vector store
            if vectorstore_type.lower() == "chroma":
                # Create a temporary directory for Chroma
                persist_directory = tempfile.mkdtemp()
                vectorstore = Chroma.from_texts(
                    texts=text_chunks, 
                    embedding=embeddings,
                    persist_directory=persist_directory
                )
            else: 
                vectorstore = FAISS.from_texts(texts=text_chunks, embedding=embeddings)
            
            # Store chunks and embeddings info for display
            st.session_state.text_chunks = text_chunks
            st.session_state.embeddings_model = embeddings
            st.session_state.vectorstore_type = vectorstore_type
            
            # Calculate a sample embedding for visualization
            sample_embedding = embeddings.embed_query(text_chunks[0][:500])
            st.session_state.sample_embedding = sample_embedding
            
            return vectorstore
    except Exception as e:
        st.error(f"Error creating vector store: {str(e)}")
        return None

def create_llm():
    """Create LLM based on selected provider and model."""
    provider = st.session_state.get("llm_provider", "gemini")
    
    if provider == "openai":
        # Use OpenAI models if selected...............
        try:
            return ChatOpenAI(
                model="gpt-3.5-turbo",
                temperature=st.session_state.get("temperature", 0.7),
                openai_api_key=st.secrets.get("OPENAI_API_KEY", "")
            )
        except Exception as e:
            st.error(f"Error initializing OpenAI: {str(e)}")
           
            st.warning("Falling back to Google Gemini model")
            provider = "gemini"
    
    # Default to Google Gemini models.........................
    try:
        return ChatGoogleGenerativeAI(
            model=st.session_state.get("gemini_model", "gemini-1.5-flash-latest"),
            temperature=st.session_state.get("temperature", 0.7),
            google_api_key=st.secrets["AIzaSyAjO68kgK4pai_vLm4qAp_L6obejrDj3Lk"]
        )
    except Exception as e:
        st.error(f"Error initializing Gemini: {str(e)}")
        return None

def get_conversation_chain(vectorstore, use_compression=False):
    if vectorstore is None:
        return None
    try:
        llm = create_llm()
        if llm is None:
            return None
            
        memory = ConversationBufferMemory(
            memory_key='chat_history',
            return_messages=True,
            output_key='answer'
        )
        
        retriever = vectorstore.as_retriever(
            search_kwargs={'k': st.session_state.get("retrieval_k", 3)}
        )
        
      
        if use_compression:
            compressor = LLMChainExtractor.from_llm(llm)
            retriever = ContextualCompressionRetriever(
                base_retriever=retriever,
                base_compressor=compressor
            )
        
        conversation_chain = ConversationalRetrievalChain.from_llm(
            llm=llm,
            retriever=retriever,
            memory=memory,
            return_source_documents=True,
            chain_type=st.session_state.get("chain_type", "stuff"),
            verbose=True
        )
        
        # Store retriever settings...
        st.session_state.retriever_settings = {
            "k": st.session_state.get("retrieval_k", 3),
            "compression": use_compression,
            "chain_type": st.session_state.get("chain_type", "stuff")
        }
        
        return conversation_chain
    except Exception as e:
        st.error(f"Error creating conversation chain: {str(e)}")
        return None

def handle_userinput(user_question):
    if not st.session_state.get("vector_store_created", False):
        st.error("Please upload and process documents first!", icon="🚨")
        return
    
    if st.session_state.conversation is None:
        st.error("Conversation chain not initialized properly", icon="🚨")
        return

    try:
      
        st.write(user_template.replace("{{MSG}}", user_question), unsafe_allow_html=True)
        
        # Add a placeholder for the bot response with a spinner
        response_placeholder = st.empty()
        response_placeholder.info("Thinking...")
        
        # Record start time for performance tracking................................................................
        start_time = time.time()
       
        response = st.session_state.conversation.invoke({
            "question": user_question
        })
        
    
        response_time = time.time() - start_time
        if "response_times" not in st.session_state:
            st.session_state.response_times = []
        st.session_state.response_times.append(response_time)
        
        # Extract answer and source docs.........................................................................
        answer = response['answer']
        source_docs = response['source_documents']
        
        response_placeholder.empty()
        st.write(bot_template.replace("{{MSG}}", answer), unsafe_allow_html=True)
        
    
        if "conversation_history" not in st.session_state:
            st.session_state.conversation_history = []
        st.session_state.conversation_history.append(
            {"role": "user", "content": user_question}
        )
        st.session_state.conversation_history.append(
            {"role": "assistant", "content": answer}
        )
        
        # Display response timeing .........................................................................
        st.caption(f"Response time: {response_time:.2f} seconds")
        
       
        col1, col2, col3 = st.columns([1, 1, 5])
        with col1:
            if st.button("👍 Helpful", key=f"helpful_{len(st.session_state.conversation_history)}"):
                if "feedback" not in st.session_state:
                    st.session_state.feedback = []
                st.session_state.feedback.append({"question": user_question, "response": answer, "rating": "helpful"})
                st.success("Thanks for your feedback!")
        
        with col2:
            if st.button("👎 Not Helpful", key=f"not_helpful_{len(st.session_state.conversation_history)}"):
                if "feedback" not in st.session_state:
                    st.session_state.feedback = []
                st.session_state.feedback.append({"question": user_question, "response": answer, "rating": "not_helpful"})
                st.error("Thanks for your feedback!")
                
       
        with st.expander("View Source Chunks"):
            for i, doc in enumerate(source_docs):
                st.markdown(f"**Relevant Chunk {i+1}:**")
                st.write(doc.page_content)
                st.markdown("---")
                
    except Exception as e:
        st.error(f"Error processing question: {str(e)}")

def generate_document_summary():
    """Generate an executive summary of the document content."""
    if not st.session_state.get("all_text", ""):
        st.error("No document content available for summarization")
        return

    try:
        llm = create_llm()
        if llm is None:
            st.error("Could not initialize LLM for summarization")
            return
            
        with st.spinner("Generating document summary..."):
            # For long documents, use a summary of first few chunks
            text_to_summarize = st.session_state.text_chunks[:min(5, len(st.session_state.text_chunks))]
            text_to_summarize = "\n\n".join(text_to_summarize)
            
            # Call the LLM for summarization
            response = llm.invoke(
                [HumanMessage(content=f"""
                Please provide a concise executive summary of the following document content.
                Focus on key points, main topics, and important findings.
                Format your response in markdown with appropriate headings.
                
                DOCUMENT CONTENT:
                {text_to_summarize}
                """)]
            )
            
            summary = response.content
            
            # Store the summary
            st.session_state.document_summary = summary
            return summary
            
    except Exception as e:
        st.error(f"Error generating summary: {str(e)}")
        return None

def display_pdf_content():
    if "all_pdf_texts" in st.session_state and st.session_state.all_pdf_texts:
        for pdf_text in st.session_state.all_pdf_texts:
            with st.expander(f"Content of {pdf_text['filename']}"):
                st.text_area("Document Content", value=pdf_text['content'], height=300)
                
                # Add download link for the extracted text
                download_link = generate_download_link(pdf_text['content'], f"{pdf_text['filename']}_extracted.txt", "Download Extracted Text")
                st.markdown(download_link, unsafe_allow_html=True)

def generate_download_link(content, filename, link_text):
    """Generate a download link for text content."""
    b64 = base64.b64encode(content.encode()).decode()
    href = f'<a href="data:file/txt;base64,{b64}" download="{filename}" class="pdf-download">{link_text}</a>'
    return href

def display_vector_data():
    if "text_chunks" in st.session_state and "embeddings_model" in st.session_state:
        chunks = st.session_state.text_chunks
        embeddings_model = st.session_state.embeddings_model
        
        # Display chunks
        with st.expander("Text Chunks"):
            chunk_limit = min(len(chunks), 20)  # Limit to prevent UI slowdown
            for i, chunk in enumerate(chunks[:chunk_limit]):
                st.markdown(f"**Chunk {i+1}/{len(chunks)}:**")
                st.text_area(f"Chunk content {i+1}", value=chunk, height=150, key=f"chunk_{i}")
                st.markdown("---")
            
            if len(chunks) > chunk_limit:
                st.info(f"Showing {chunk_limit} of {len(chunks)} chunks to maintain performance.")
        
        # Display chunk statistics
        with st.expander("Chunk Statistics"):
            # Calculate and display statistics about chunks
            chunk_lengths = [len(chunk) for chunk in chunks]
            
            st.write(f"**Total chunks:** {len(chunks)}")
            st.write(f"**Average chunk length:** {sum(chunk_lengths)/len(chunk_lengths):.2f} characters")
            st.write(f"**Min chunk length:** {min(chunk_lengths)} characters")
            st.write(f"**Max chunk length:** {max(chunk_lengths)} characters")
            
            # Create histogram of chunk lengths
            fig, ax = plt.subplots()
            ax.hist(chunk_lengths, bins=20)
            ax.set_title('Distribution of Chunk Lengths')
            ax.set_xlabel('Chunk Length (characters)')
            ax.set_ylabel('Frequency')
            st.pyplot(fig)
        
        # Display embeddings info
        with st.expander("Vector Embeddings"):
            st.write(f"**Embedding model:** {embeddings_model.model_name}")
            st.write(f"**Vector store type:** {st.session_state.get('vectorstore_type', 'FAISS')}")
            
            # Show sample embedding
            if "sample_embedding" in st.session_state:
                sample = st.session_state.sample_embedding
                st.write(f"**Embedding dimensions:** {len(sample)}")
                
                # Create interactive visualization with Plotly
                fig = px.line(
                    x=list(range(len(sample))), 
                    y=sample[:100],  # Show first 100 dimensions only
                    title="Sample Embedding Vector (First 100 Dimensions)",
                    labels={"x": "Dimension", "y": "Value"}
                )
                st.plotly_chart(fig, use_container_width=True)

def export_conversation():
    """Export the conversation history to a downloadable format."""
    if "conversation_history" not in st.session_state or not st.session_state.conversation_history:
        st.warning("No conversation to export")
        return
        
    try:
        # Format the conversation as markdown
        markdown_content = "# Document Chat Conversation\n\n"
        markdown_content += f"*Generated on: {time.strftime('%Y-%m-%d %H:%M:%S')}*\n\n"
        
     
        if "file_names" in st.session_state:
            markdown_content += "## Documents\n"
            for file in st.session_state.file_names:
                markdown_content += f"- {file}\n"
            markdown_content += "\n"
        
      
        markdown_content += "## Conversation\n\n"
        for entry in st.session_state.conversation_history:
            role = "**User**" if entry["role"] == "user" else "**AI Assistant**"
            markdown_content += f"{role}:\n\n{entry['content']}\n\n---\n\n"
        
      
        b64 = base64.b64encode(markdown_content.encode()).decode()
        date_str = time.strftime('%Y%m%d_%H%M%S')
        filename = f"conversation_export_{date_str}.md"
        href = f'<a href="data:text/markdown;base64,{b64}" download="{filename}" class="pdf-download">Download Conversation</a>'
        
        st.markdown(href, unsafe_allow_html=True)
        
    except Exception as e:
        st.error(f"Error exporting conversation: {str(e)}")

def display_analytics():
    """Display analytics about the conversation and system performance."""
    # Check if we have any data to display
    if "response_times" not in st.session_state or not st.session_state.response_times:
        st.info("No analytics data available yet. Try asking some questions first.")
        return
        
    # Response time analytics
    st.subheader("Response Time Analytics")
    response_times = st.session_state.response_times
    
    # Calculate statistics
    avg_time = sum(response_times) / len(response_times)
    max_time = max(response_times)
    min_time = min(response_times)
    
    # Display metrics
    col1, col2, col3 = st.columns(3)
    col1.metric("Average Response Time", f"{avg_time:.2f}s")
    col2.metric("Fastest Response", f"{min_time:.2f}s")
    col3.metric("Slowest Response", f"{max_time:.2f}s")
    
    # Create a line chart of response times
    times_df = pd.DataFrame({
        "Query": range(1, len(response_times) + 1),
        "Response Time (s)": response_times
    })
    st.line_chart(times_df.set_index("Query"))
    
    # Conversation statistics
    if "conversation_history" in st.session_state:
        st.subheader("Conversation Statistics")
        history = st.session_state.conversation_history
        
        # Count messages by role
        user_messages = sum(1 for msg in history if msg["role"] == "user")
        ai_messages = sum(1 for msg in history if msg["role"] == "assistant")
        
        # Calculate average message length
        user_lengths = [len(msg["content"]) for msg in history if msg["role"] == "user"]
        ai_lengths = [len(msg["content"]) for msg in history if msg["role"] == "assistant"]
        
        avg_user_length = sum(user_lengths) / len(user_lengths) if user_lengths else 0
        avg_ai_length = sum(ai_lengths) / len(ai_lengths) if ai_lengths else 0
        
        # Display metrics
        col1, col2 = st.columns(2)
        col1.metric("User Messages", user_messages)
        col2.metric("AI Responses", ai_messages)
        
        col1, col2 = st.columns(2)
        col1.metric("Avg. User Message Length", f"{avg_user_length:.0f} chars")
        col2.metric("Avg. AI Response Length", f"{avg_ai_length:.0f} chars")
    
    # Show feedback statistics if available
    if "feedback" in st.session_state and st.session_state.feedback:
        st.subheader("Feedback Statistics")
        feedback = st.session_state.feedback
        
        # Count helpful vs. not helpful
        helpful = sum(1 for item in feedback if item["rating"] == "helpful")
        not_helpful = sum(1 for item in feedback if item["rating"] == "not_helpful")
        
        # Create pie chart
        labels = ['Helpful', 'Not Helpful']
        values = [helpful, not_helpful]
        
        fig = px.pie(values=values, names=labels, title="Response Feedback Distribution")
        st.plotly_chart(fig, use_container_width=True)

def settings_page():
    """Display and manage application settings."""
    st.subheader("Application Settings")
    
    # Create sections with expanders
    with st.expander("LLM Settings", expanded=True):
        # LLM Provider
        provider = st.selectbox(
            "LLM Provider",
            options=["gemini", "openai"],
            index=0 if st.session_state.get("llm_provider", "gemini") == "gemini" else 1,
            key="settings_llm_provider"
        )
        
        # Model selection based on provider
        if provider == "gemini":
            model = st.selectbox(
                "Gemini Model",
                options=["gemini-1.5-flash-latest", "gemini-1.5-pro-latest"],
                index=0 if st.session_state.get("gemini_model", "gemini-1.5-flash-latest") == "gemini-1.5-flash-latest" else 1,
                key="settings_gemini_model"
            )
            st.session_state.gemini_model = model
        else:
            # OpenAI models would go here
            st.info("Using OpenAI gpt-3.5-turbo model")
        
        # Temperature
        temperature = st.slider(
            "Temperature",
            min_value=0.0,
            max_value=1.0,
            value=st.session_state.get("temperature", 0.7),
            step=0.1,
            key="settings_temperature"
        )
        
        # Store LLM settings
        st.session_state.llm_provider = provider
        st.session_state.temperature = temperature
    
    with st.expander("Vector Store Settings"):
        # Embedding model
        embedding_model = st.selectbox(
            "Embedding Model",
            options=["all-MiniLM-L6-v2", "all-mpnet-base-v2"],
            index=0 if st.session_state.get("embedding_model", "all-MiniLM-L6-v2") == "all-MiniLM-L6-v2" else 1,
            key="settings_embedding_model"
        )
        st.session_state.embedding_model = embedding_model
        
        # Vector store type
        vectorstore_type = st.selectbox(
            "Vector Store Type",
            options=["faiss", "chroma"],
            index=0 if st.session_state.get("vectorstore_type", "faiss") == "faiss" else 1,
            key="settings_vectorstore_type"
        )
        st.session_state.vectorstore_type = vectorstore_type
        
        # Retrieval parameters
        retrieval_k = st.slider(
            "Number of chunks to retrieve (k)",
            min_value=1,
            max_value=10,
            value=st.session_state.get("retrieval_k", 3),
            step=1,
            key="settings_retrieval_k"
        )
        st.session_state.retrieval_k = retrieval_k
        
        # Chain type
        chain_type = st.selectbox(
            "Chain Type",
            options=["stuff", "map_reduce", "refine"],
            index=0 if st.session_state.get("chain_type", "stuff") == "stuff" else 
                  (1 if st.session_state.get("chain_type") == "map_reduce" else 2),
            key="settings_chain_type"
        )
        st.session_state.chain_type = chain_type
        
        # Compression
        use_compression = st.checkbox(
            "Use contextual compression",
            value=st.session_state.get("use_compression", False),
            key="settings_use_compression"
        )
        st.session_state.use_compression = use_compression
    
    with st.expander("Text Processing Settings"):
        # Chunking method
        chunking_method = st.selectbox(
            "Text Chunking Method",
            options=["recursive", "character"],
            index=0 if st.session_state.get("chunking_method", "recursive") == "recursive" else 1,
            key="settings_chunking_method"
        )
        st.session_state.chunking_method = chunking_method
        
        # Chunk size
        chunk_size = st.slider(
            "Chunk Size",
            min_value=500,
            max_value=2000,
            value=st.session_state.get("chunk_size", 1000),
            step=100,
            key="settings_chunk_size"
        )
        st.session_state.chunk_size = chunk_size
        
        # Chunk overlap
        chunk_overlap = st.slider(
            "Chunk Overlap",
            min_value=0,
            max_value=500,
            value=st.session_state.get("chunk_overlap", 200),
            step=50,
            key="settings_chunk_overlap"
        )
        st.session_state.chunk_overlap = chunk_overlap
    
    # Save settings button
    if st.button("Save Settings", key="save_settings"):
        st.success("Settings saved successfully!")
        # Reset conversation if settings changed
        if st.session_state.get("vector_store_created", False):
            st.warning("Note: Some settings will only take effect when you re-process your documents.")

def reset_conversation():
    """Reset the conversation history."""
    if "conversation_history" in st.session_state:
        del st.session_state.conversation_history
    if "response_times" in st.session_state:
        del st.session_state.response_times
    return True

def main():
    st.set_page_config(page_title="Enhanced PDF Chat", page_icon="📚", layout="wide")
    st.write(css, unsafe_allow_html=True)

    # Initialize session state
    if "conversation" not in st.session_state:
        st.session_state.conversation = None
    if "vector_store_created" not in st.session_state:
        st.session_state.vector_store_created = False
    if "file_names" not in st.session_state:
        st.session_state.file_names = []
    if "embedding_model" not in st.session_state:
        st.session_state.embedding_model = "all-MiniLM-L6-v2"
    if "llm_provider" not in st.session_state:
        st.session_state.llm_provider = "gemini"
    if "gemini_model" not in st.session_state:
        st.session_state.gemini_model = "gemini-1.5-flash-latest"
    if "temperature" not in st.session_state:
        st.session_state.temperature = 0.7
    if "chunking_method" not in st.session_state:
        st.session_state.chunking_method = "recursive"
        
    # Custom sidebar
    with st.sidebar:
        st.title("📚 Enhanced PDF Chat")
        st.subheader("Document Processing")
        
        # File uploader with multiple file types
        uploaded_files = st.file_uploader(
            "Upload your documents",
            accept_multiple_files=True,
            type=['pdf', 'docx', 'txt', 'zip']
        )
        
        if uploaded_files:
            st.session_state.file_names = [file.name for file in uploaded_files]
        
        # Advanced processing options
        with st.expander("Processing Options"):
            # Use settings from session state
            chunking_method = st.radio(
                "Chunking Method",
                options=["recursive", "character"],
                index=0 if st.session_state.chunking_method == "recursive" else 1
            )
            
            vectorstore_type = st.radio(
                "Vector Store Type",
                options=["faiss", "chroma"],
                index=0 if st.session_state.get("vectorstore_type", "faiss") == "faiss" else 1
            )
            
            use_compression = st.checkbox(
                "Use contextual compression", 
                value=st.session_state.get("use_compression", False)
            )
        
        # Process button
        if st.button("Process Documents"):
            if not uploaded_files:
                st.error("Please upload at least one document.")
            else:
                with st.spinner("Processing your documents..."):
                    # Reset the state
                    st.session_state.vector_store_created = False
                    st.session_state.conversation = None
                    
                    # Process each file
                    all_text = ""
                    all_texts = []
                    
                    for file in uploaded_files:
                        st.info(f"Processing {file.name}...")
                        file_text = parse_file(file)
                        
                        if file_text.strip():
                            all_text += file_text + "\n\n"
                            all_texts.append({
                                "filename": file.name,
                                "content": file_text
                            })
                    
                    if all_text.strip():
                        # Store the raw text
                        st.session_state.all_text = all_text
                        st.session_state.all_pdf_texts = all_texts
                        
                        # Process the text
                        text_chunks = get_text_chunks(all_text, method=chunking_method)
                        if text_chunks:
                            vectorstore = get_vectorstore(text_chunks, vectorstore_type=vectorstore_type)
                            if vectorstore:
                                conversation_chain = get_conversation_chain(vectorstore, use_compression=use_compression)
                                if conversation_chain:
                                    st.session_state.conversation = conversation_chain
                                    st.session_state.vector_store_created = True
                                    st.success("Processing complete! You can now chat with your documents.")
                                    
                                    # Generate summary
                                    with st.spinner("Generating document summary..."):
                                        summary = generate_document_summary()
                                        if summary:
                                            st.success("Summary generated successfully!")
                    else:
                        st.error("No text could be extracted from the documents.")
        
        # Additional sidebar options
        if st.session_state.get("vector_store_created", False):
            if st.button("Reset Conversation"):
                if reset_conversation():
                    st.success("Conversation reset successfully!")
            
            if st.button("Export Conversation"):
                export_conversation()
                
        # Show app information
        st.markdown("---")
        st.markdown("### About Enhanced PDF Chat")
        st.markdown("""
        This application allows you to chat with your documents using AI. 
        Features include:
        - Multi-format document support (PDF, DOCX, TXT, ZIP)
        - Advanced OCR for scanned documents
        - Multiple AI models support
        - Analytics and conversation export
        """)

    # Create main tabs
    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "💬 Chat", 
        "📄 Document Content", 
        "🔍 Vector Data", 
        "📊 Analytics", 
        "⚙️ Settings"
    ])
    
    with tab1:
       
        st.header("Chat with Your Documents")
        
        # Display document summary if available
        if "document_summary" in st.session_state:
            with st.expander("Document Summary", expanded=True):
                st.markdown(st.session_state.document_summary)
        
        # Chat history container
        chat_container = st.container()
        
        with chat_container:
            
            if "conversation_history" in st.session_state:
                for message in st.session_state.conversation_history:
                    if message["role"] == "user":
                        st.write(user_template.replace("{{MSG}}", message["content"]), unsafe_allow_html=True)
                    else:
                        st.write(bot_template.replace("{{MSG}}", message["content"]), unsafe_allow_html=True)
        
        # Input for user question......................
        st.markdown("---")
        user_question = st.text_input(
            "Ask a question about your documents:",
            placeholder="What would you like to know about your documents?",
            key="user_question"
        )
        
        # Add send button....................
        col1, col2 = st.columns([6, 1])
        with col2:
            send_pressed = st.button("Send", key="send_question")
            
        if user_question and send_pressed:
            handle_userinput(user_question)
            # Clear the input field after sending...............
            st.session_state.user_question = ""
    
    with tab2:
        # Document Content display...................
        st.header("Document Content")
        display_pdf_content()
    
    with tab3:
        # Vector data display.........
        st.header("Vector Data")
        display_vector_data()
    
    with tab4:
        # Analytics display.............
        st.header("Analytics")
        display_analytics()
    
    with tab5:
        # Settings page...............
        settings_page()

if __name__ == '__main__':
    main()